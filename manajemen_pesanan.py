import docx
from docx.shared import Inches
import datetime
import os
import openpyxl


def get_harga_menu(menu):
    harga_menu = {"Nasi Goreng": 20000, "Bakso": 15000, "Mie Ayam": 12000}
    return harga_menu.get(menu.strip(), 0)


def save_to_docx(pesanan):
    doc = docx.Document()
    doc.add_heading('Detail Pesanan', 0)
    for key, value in pesanan.items():
        doc.add_paragraph(f"{key}: {value}")

    if pesanan['image_path']:
        try:
            doc.add_picture(pesanan['image_path'], width=Inches(2))
        except Exception as e:
            print(f"Error menambahkan gambar: {e}")

    doc.save(f"pesanan_{pesanan['nama_pelanggan']}.docx")


def save_to_xlsx(pesanan):
    file_exists = os.path.isfile('pesanan.xlsx')
    if file_exists:
        workbook = openpyxl.load_workbook('pesanan.xlsx')
        sheet = workbook.active
    else:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        header = list(pesanan.keys())
        sheet.append(header)

    sheet.append(list(pesanan.values()))
    workbook.save('pesanan.xlsx')


def create_pesanan():
    nama_pelanggan = input("Nama Pelanggan: ").strip()
    if not nama_pelanggan:
        print("Nama pelanggan tidak boleh kosong.")
        return

    menu = input("Menu (pisahkan dengan koma): ").split(",")
    if not menu:
        print("Menu tidak boleh kosong.")
        return

    try:
        jumlah = [int(x) for x in input("Jumlah : ").split(",")]
    except ValueError:
        print("Jumlah harus berupa angka.")
        return

    catatan = input("Catatan Tambahan: ")
    image_path = input("Path Gambar: ")

    pesanan = {
        "nama_pelanggan": nama_pelanggan,
        "menu": ','.join(menu),
        "jumlah": ','.join(map(str, jumlah)),
        "status": "Sedang Diproses",
        "tanggal": datetime.datetime.now().strftime("%Y-%m-%d"),
        "waktu": datetime.datetime.now().strftime("%H:%M:%S"),
        "catatan": catatan,
        "image_path": image_path
    }

    harga_total = sum(get_harga_menu(m) * j for m, j in zip(menu, jumlah))
    pesanan['harga_total'] = harga_total

    save_to_docx(pesanan)
    save_to_xlsx(pesanan)

    print("Pesanan berhasil dibuat!")



def read_pesanan():
    try:
        workbook = openpyxl.load_workbook('pesanan.xlsx')
        sheet = workbook.active
        headers = [cell.value for cell in sheet[1]]  
        index_nama = headers.index('nama_pelanggan')   
        index_status = headers.index('status')   
        index_menu = headers.index('menu')  
        index_jumlah = headers.index('jumlah') 

        print("\nPesanan yang ada:")
        for row in sheet.iter_rows(min_row=2, values_only=True):
            row_list = list(row) 
            print(f"\nNama : {row_list[index_nama]}")
            print(f"Menu : {row_list[index_menu]}")
            print(f"Jumlah : {row_list[index_jumlah]}")
            print(f"Status : {row_list[index_status]}")

    except FileNotFoundError:
        print("File pesanan.xlsx tidak ditemukan.")




def update_pesanan():
    nama_pelanggan = input("Masukkan Nama Pelanggan yang ingin diperbarui: ").strip().lower()
    found = False
    updated_rows = []

    try:
        workbook = openpyxl.load_workbook('pesanan.xlsx')
        sheet = workbook.active
        headers = [cell.value for cell in sheet[1]]  
        index_nama = headers.index('nama_pelanggan')  
        index_status = headers.index('status') 
        index_menu = headers.index('menu') 
        index_jumlah = headers.index('jumlah') 

        for row in sheet.iter_rows(min_row=2, values_only=True):
            row_list = list(row)  
            if row_list[index_nama].strip().lower() == nama_pelanggan:
                found = True
                print(f"\nPesanan ditemukan:")
                print(f"Nama : {row_list[index_nama]}")
                print(f"Menu : {row_list[index_menu]}")
                print(f"Jumlah : {row_list[index_jumlah]}")
                print(f"Status : {row_list[index_status]}")

                status_baru = input("Masukkan status baru (Proses/Selesai/Batal): ").strip().capitalize()

                if status_baru in ['Proses', 'Selesai', 'Batal']:
                    row_list[index_status] = status_baru
                    print(f"\nStatus berhasil diperbarui menjadi: {status_baru}")
                else:
                    print("Status tidak valid. Status tetap tidak berubah.")

            updated_rows.append(row_list)

        if found:
            sheet.delete_rows(2, sheet.max_row + 1) 
            sheet.delete_rows(2)
            for row in updated_rows:
                sheet.append(row)  
            workbook.save('pesanan.xlsx')
            print("Pesanan berhasil diperbarui!")
        else:
            print("Pesanan dengan nama tersebut tidak ditemukan.")

    except FileNotFoundError:
        print("File pesanan.xlsx tidak ditemukan.")


def delete_pesanan():
    nama_pelanggan = input("Masukkan Nama Pelanggan yang ingin dihapus: ").strip()
    rows_to_keep = []
    found = False

    try:
        workbook = openpyxl.load_workbook('pesanan.xlsx')
        sheet = workbook.active

        for row in sheet.iter_rows(min_row=2, values_only=True):
            if row[0].strip().lower() != nama_pelanggan.lower():
                rows_to_keep.append(row)
            else:
                found = True

        if found:
            for i in range(2, sheet.max_row + 1):
                sheet.delete_rows(2)
            for row in rows_to_keep:
                sheet.append(row)
            workbook.save('pesanan.xlsx')
            print("Pesanan berhasil dihapus!")
        else:
            print("Pesanan dengan nama tersebut tidak ditemukan.")

    except FileNotFoundError:
        print("File pesanan.xlsx tidak ditemukan.")

def search_pesanan():
    nama_pelanggan = input("Masukkan Nama Pelanggan yang ingin dicari: ").strip().lower()
    found = False

    try:
        workbook = openpyxl.load_workbook('pesanan.xlsx')
        sheet = workbook.active
        headers = [cell.value for cell in sheet[1]]  
        index_nama = headers.index('nama_pelanggan')  
        index_status = headers.index('status')  
        index_menu = headers.index('menu')  
        index_jumlah = headers.index('jumlah')  

        print("\nHasil Pencarian:")
        for row in sheet.iter_rows(min_row=2, values_only=True):
            row_list = list(row) 
            if row_list[index_nama].strip().lower() == nama_pelanggan:
                found = True
                print(f"\nNama : {row_list[index_nama]}")
                print(f"Menu : {row_list[index_menu]}")
                print(f"Jumlah : {row_list[index_jumlah]}")
                print(f"Status : {row_list[index_status]}")
                break  

        if not found:
            print("Pesanan dengan nama tersebut tidak ditemukan.")

    except FileNotFoundError:
        print("File pesanan.xlsx tidak ditemukan.")


while True:
    print("\nMenu Manajemen Pesanan Restoran")
    print("1. Buat Pesanan")
    print("2. Baca Pesanan")
    print("3. Perbarui Pesanan")
    print("4. Hapus Pesanan")
    print("5. Cari Pesanan")
    print("6. Keluar")
    pilihan = input("Pilih opsi (1-6): ")

    if pilihan == '1':
        create_pesanan()
    elif pilihan == '2':
        read_pesanan()
    elif pilihan == '3':
        update_pesanan()
    elif pilihan == '4':
        delete_pesanan()
    elif pilihan == '5':
        search_pesanan()
    elif pilihan == '6':
        print("Keluar dari program.")
        break
    else:
        print("Pilihan tidak valid. Silakan coba lagi.")
